import os
from collections import defaultdict
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class MusicLibrary:
    def __init__(self):
        self.supported_formats = ('.mp3', '.wav', '.ogg', '.flac', '.m4a', '.aac', '.wma')
        self.music_library = defaultdict(lambda: defaultdict(dict))

    def scan_folder(self, folder_path, clear_existing=True):
        """
        Сканирует папку и добавляет музыку в коллекцию
        :param folder_path: Путь к сканируемой папке
        :param clear_existing: Если True, очищает существующую коллекцию перед сканированием
        :return: True если успешно, False если отменено
        """
        if not folder_path:
            return False

        self.music_library.clear()
        self.music_library = defaultdict(lambda: defaultdict(dict))  # Восстанавливаем структуру

        self._scan_folder_recursive(folder_path, self.music_library)
        return True

    def _scan_folder_recursive(self, current_path, node):
        for entry in os.listdir(current_path):
            full_path = os.path.join(current_path, entry)

            if os.path.isdir(full_path):
                if entry not in node:
                    node[entry] = {}
                self._scan_folder_recursive(full_path, node[entry])
            elif entry.lower().endswith(self.supported_formats):
                if "_files" not in node:
                    node["_files"] = []
                node["_files"].append((entry, full_path))

    def get_library(self):
        return self.music_library

    def clear_library(self):
        """Полностью очищает музыкальную коллекцию"""
        self.music_library.clear()
        # Восстанавливаем структуру defaultdict
        self.music_library = defaultdict(lambda: defaultdict(dict))

    # def save_to_docx(self, output_path=None):
    #     if not output_path:
    #         output_path = filedialog.asksaveasfilename(
    #             defaultextension=".docx",
    #             filetypes=[("Word Documents", "*.docx")],
    #             title="Сохранить коллекцию как"
    #         )
    #         if not output_path:
    #             return False, ""
    #
    #     try:
    #         doc = Document()
    #         style = doc.styles['Normal']
    #         style.font.name = 'Arial'
    #         style.font.size = Pt(12)
    #
    #         # Заголовок документа
    #         title = doc.add_heading('Моя музыкальная коллекция', level=1)
    #         title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    #         def add_items(node, level=1):
    #             for name, content in node.items():
    #                 if name == "_files":
    #                     for file_name, _ in content:
    #                         p = doc.add_paragraph('    ' * level + f"🎵 {file_name}")
    #                         p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    #                 else:
    #                     heading = doc.add_heading('    ' * (level - 1) + f"📁 {name}", level=min(level + 1, 6))
    #                     heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)
    #                     add_items(content, level + 1)
    #
    #         add_items(self.music_library)
    #         doc.save(output_path)
    #         return True, output_path
    #
    #     except Exception as e:
    #         return False, str(e)

    def save_to_docx(self, output_path=None):
        if not output_path:
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Сохранить коллекцию как"
            )
            if not output_path:
                return False, ""

        try:
            doc = Document()

            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)

            # Создаем стиль для папок
            if 'FolderStyle' not in doc.styles:
                folder_style = doc.styles.add_style('FolderStyle', WD_STYLE_TYPE.PARAGRAPH)
                folder_style.base_style = doc.styles['Heading 2']
                folder_style.font.color.rgb = RGBColor(0, 0, 128)
                folder_style.font.bold = True

            # Заголовок документа
            title = doc.add_heading('Моя музыкальная коллекция', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Рекурсивное добавление структуры
            def add_items(node, level=1):
                for name, content in node.items():
                    if name == "_files":
                        continue

                    # Добавляем папку с полем COLLAPSIBLE
                    para = doc.add_paragraph(style='FolderStyle')
                    run = para.add_run(f"📁 {name}")

                    # Добавляем поле сворачивания
                    doc.paragraphs[-1]._element.append(
                        self._create_collapsible_field(f"folder_{name}"))

                    # Добавляем файлы (будут скрыты в свёрнутом состоянии)
                    if "_files" in node:
                        for file_name, _ in sorted(node["_files"], key=lambda x: x[0]):
                    file_para = doc.add_paragraph(f"    🎵 {file_name}", style='List Bullet')
                    file_para.paragraph_format.left_indent = Pt(level * 20)

                    # Рекурсивно добавляем подпапки
                    add_items(content, level + 1)

                    add_items(self.music_library)

                    # Включаем обновление полей при открытии
                    doc.settings.update_fields_on_open = True
                    doc.save(output_path)

            return True, output_path

        except Exception as e:
            return False, str(e)

    def _create_collapsible_field(self, bookmark_name):
        """Создает XML-элемент для сворачиваемого поля"""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Создаем элемент поля
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = f' COLLAPSIBLE {{}} \\b "{bookmark_name}"'

        end_fld_char = OxmlElement('w:fldChar')
        end_fld_char.set(qn('w:fldCharType'), 'end')

        # Собираем все элементы вместе
        r = OxmlElement('w:r')
        r.append(fld_char)
        r.append(instr_text)
        r.append(end_fld_char)

        return r