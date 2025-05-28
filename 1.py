import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from foobar_controller import FoobarController

class MusicCollectionApp:
    def __init__(self, root):
        """Инициализация главного окна приложения"""
        self.root = root
        self.root.title("Твоя Музыка")  # Заголовок окна
        self.root.geometry("1000x700")  # Размер окна

        # Конфигурация путей и форматов
        self.foobar_path = r"C:\Program Files (x86)\foobar2000\foobar2000.exe"  # Путь к foobar2000
        self.supported_formats = ('.mp3', '.wav', '.ogg', '.flac', '.m4a', '.aac', '.wma')  # Поддерживаемые форматы

        # Структура данных для хранения коллекции
        self.music_library = defaultdict(lambda: defaultdict(dict))

        # Инициализация интерфейса
        self._setup_ui()
# ====================================
        # Инициализация контроллера
        self.foobar = FoobarController()
#====================================
    def _setup_ui(self):
        """Создание пользовательского интерфейса"""
        # Верхняя панель с заголовком
        header_frame = tk.Frame(self.root, bg="#f0f0f0", padx=10, pady=10)
        header_frame.pack(fill=tk.X)
        tk.Label(header_frame,
                 text="ТВОЯ МУЗЫКА",
                 font=('Arial', 14, 'bold'),
                 bg="#f0f0f0").pack(side=tk.LEFT)

        # Панель инструментов
        toolbar = tk.Frame(self.root, padx=5, pady=5)
        toolbar.pack(fill=tk.X)

        # Кнопки управления
        buttons = [
            ("📁 Сканировать папку", self.scan_folder),
            ("🗑️ Очистить коллекцию", self.clear_library)
        ]
        # 28.05.25 Добавьте кнопку в интерфейс (в метод _setup_ui):
        #=============================================
        btn_export = tk.Button(toolbar,
                               text="💾 Экспорт в DOCX",
                               command=self.export_to_docx,
                               bd=1,
                               relief=tk.RIDGE,
                               padx=10)
        btn_export.pack(side=tk.RIGHT, padx=2)
        #=================================================
        for text, cmd in buttons:
            btn = tk.Button(toolbar, text=text, command=cmd, bd=1, relief=tk.RIDGE, padx=10)
            btn.pack(side=tk.LEFT, padx=2)

        # Дерево коллекции
        tree_frame = tk.Frame(self.root)
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.tree = ttk.Treeview(tree_frame, columns=("type", "path"), show="tree", selectmode="extended")

        # Настройка стилей для дерева
        self.tree.tag_configure("folder", background="#f0f0f0", font=('Arial', 10, 'bold'))
        self.tree.tag_configure("file", background="white")

        # Scrollbar для дерева
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=scrollbar.set)
        self.tree.pack(fill=tk.BOTH, expand=True)

        # Панель управления плеером
        player_frame = tk.Frame(self.root, bg="#e0e0e0", padx=10, pady=8)
        player_frame.pack(fill=tk.X)

        # Кнопки управления плеером
        player_controls = [
            ("▶ Воспроизвести", self.play_selected),
            ("⏏ В плейлист", self.add_to_playlist),
            ("⏹ Стоп", self.stop_foobar)
        ]

        for text, cmd in player_controls:
            btn = tk.Button(player_frame, text=text, command=cmd, bg="#f8f8f8", padx=10)
            btn.pack(side=tk.LEFT, padx=5)

        # Строка состояния
        self.status_bar = tk.Label(player_frame,
                                   text="Готов к работе",
                                   bg="#e0e0e0",
                                   fg="#333333",
                                   anchor=tk.W)
        self.status_bar.pack(side=tk.LEFT, padx=10, expand=True, fill=tk.X)

        # Контекстное меню
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Воспроизвести", command=self.play_selected)
        self.context_menu.add_command(label="Добавить в плейлист", command=self.add_to_playlist)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Открыть в проводнике", command=self.open_in_explorer)
        self.context_menu.add_command(label="Удалить", command=self.delete_item)

        # Привязка событий
        self.tree.bind("<Button-3>", self.show_context_menu)
        self.tree.bind("<Double-1>", self.play_selected)

    def scan_folder(self):
        """Сканирование выбранной папки и добавление музыки в коллекцию"""
        folder_path = filedialog.askdirectory(title="Выберите папку с музыкой")
        if not folder_path:
            return

        self.status_bar.config(text=f"Сканирование: {folder_path}", fg="blue")
        self.root.update()  # Обновляем интерфейс

        try:
            # Рекурсивное сканирование
            self._scan_folder_recursive(folder_path, self.music_library)
            self.update_tree_view()
            self.status_bar.config(text=f"Добавлено: {folder_path}", fg="green")
        except Exception as e:
            self.status_bar.config(text=f"Ошибка: {str(e)}", fg="red")

    def _scan_folder_recursive(self, current_path, node):
        """Рекурсивное сканирование папок"""
        for entry in os.listdir(current_path):
            full_path = os.path.join(current_path, entry)

            if os.path.isdir(full_path):
                # Обработка подпапки
                if entry not in node:
                    node[entry] = {}
                self._scan_folder_recursive(full_path, node[entry])
            elif entry.lower().endswith(self.supported_formats):
                # Обработка аудиофайла
                if "_files" not in node:
                    node["_files"] = []
                node["_files"].append((entry, full_path))

    def update_tree_view(self):
        """Обновление отображения дерева коллекции"""
        self.tree.delete(*self.tree.get_children())  # Очищаем текущее отображение

        # Рекурсивное построение дерева
        self._build_tree_recursive("", self.music_library)

    def _build_tree_recursive(self, parent_id, node):
        """Рекурсивное построение элементов дерева"""
        for name, content in node.items():
            if name == "_files":
                # Добавление файлов
                for file_name, file_path in content:
                    self.tree.insert(
                        parent_id, "end",
                        text=file_name,
                        values=("file", file_path),
                        tags=("file",)
                    )
            else:
                # Добавление папки
                folder_id = self.tree.insert(
                    parent_id, "end",
                    text=name,
                    values=("folder", ""),
                    tags=("folder",)
                )
                self._build_tree_recursive(folder_id, content)

        # ==================================================================



    # 28.05.25 Добавьте метод в класс MusicCollectionApp:
    def export_to_docx(self):
        """Экспорт коллекции в DOCX файл"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Сохранить коллекцию как"
        )
        if file_path:
            if save_to_docx(self.music_library, file_path):
                self.status_bar.config(text=f"Коллекция сохранена: {file_path}", fg="green")
            else:
                self.status_bar.config(text="Ошибка при сохранении файла", fg="red")
    # ==============================================================================================

    def play_selected(self, event=None):
        """Воспроизведение выбранных треков в foobar2000"""
        selected_items = self.tree.selection()
        if not selected_items:
            return

        paths = []
        names = []

        for item in selected_items:
            item_data = self.tree.item(item)
            if item_data["values"][0] == "file":
                paths.append(item_data["values"][1])
                names.append(item_data["text"])
            elif item_data["values"][0] == "folder":
                self._collect_files_from_folder(item, paths, names)

        if paths:
            self._play_in_foobar(paths, " | ".join(names))

    def _collect_files_from_folder(self, folder_item, paths, names):
        """Сбор всех файлов из папки"""
        for child in self.tree.get_children(folder_item):
            child_data = self.tree.item(child)
            if child_data["values"][0] == "file":
                paths.append(child_data["values"][1])
                names.append(child_data["text"])
            elif child_data["values"][0] == "folder":
                self._collect_files_from_folder(child, paths, names)

    def _play_in_foobar(self, file_paths, display_name):
        try:
            self.foobar.play_files(file_paths)
            self.status_bar.config(text=f"Воспроизведение: {display_name[:50]}...", fg="blue")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def add_to_playlist(self):
        selected_items = self.tree.selection()
        if not selected_items:
            return

        paths = self._get_selected_paths(selected_items)
        if paths:
            try:
                self.foobar.add_to_playlist(paths)
                self.status_bar.config(text=f"Добавлено {len(paths)} треков в плейлист", fg="green")
            except Exception as e:
                messagebox.showerror("Ошибка", str(e))

    def stop_foobar(self):
        try:
            self.foobar.stop_playback()
            self.status_bar.config(text="Воспроизведение остановлено", fg="black")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))
    def open_in_explorer(self):
        """Открытие папки с выбранным файлом в проводнике"""
        selected = self.tree.focus()
        if not selected:
            return

        item_data = self.tree.item(selected)
        if item_data["values"][0] == "file":
            path = os.path.dirname(item_data["values"][1])
        else:
            path = self._find_folder_path(selected)

        if path and os.path.exists(path):
            os.startfile(path)

    def _find_folder_path(self, folder_item):
        """Поиск физического пути к папке в дереве"""
        parts = []
        current_item = folder_item

        while current_item:
            item_data = self.tree.item(current_item)
            parts.insert(0, item_data["text"])
            current_item = self.tree.parent(current_item)

        test_path = os.path.join(*parts)
        if os.path.exists(test_path):
            return test_path

        return None

    def delete_item(self):
        """Удаление выбранных элементов из коллекции"""
        selected_items = self.tree.selection()
        if not selected_items:
            return

        confirm = messagebox.askyesno(
            "Подтверждение",
            f"Удалить {len(selected_items)} выбранных элементов?"
        )
        if not confirm:
            return

        for item in selected_items:
            self._delete_item_recursive(item)

        self.update_tree_view()
        self.status_bar.config(text=f"Удалено {len(selected_items)} элементов", fg="orange")

    def _delete_item_recursive(self, item):
        """Рекурсивное удаление элемента из структуры данных"""
        parent = self.tree.parent(item)
        if not parent:  # Корневой элемент
            return

        item_data = self.tree.item(item)
        parent_data = self.tree.item(parent)

        # Удаление из структуры данных
        if parent == "":  # Верхний уровень
            if item_data["text"] in self.music_library:
                del self.music_library[item_data["text"]]
        else:
            # Находим соответствующий узел в структуре данных
            current_node = self.music_library
            path = []

            # Строим путь от корня к элементу
            current_item = item
            while current_item and current_item != "":
                path.insert(0, self.tree.item(current_item, "text"))
                current_item = self.tree.parent(current_item)

            # Находим родительский узел
            for step in path[:-1]:
                if step in current_node:
                    current_node = current_node[step]
                else:
                    return

            # Удаляем целевой узел
            if path[-1] in current_node:
                del current_node[path[-1]]

    def clear_library(self):
        """Полная очистка музыкальной коллекции"""
        if messagebox.askyesno("Подтверждение", "Очистить всю коллекцию?"):
            self.music_library.clear()
            self.update_tree_view()
            self.status_bar.config(text="Коллекция очищена", fg="red")

    def show_context_menu(self, event):
        """Отображение контекстного меню"""
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

# ===============================================================
# 28.05.25 функция для сохранения структуры музыкальной коллекции в файл .docx с библиотекой python-docx
def save_to_docx(music_library, output_path="music_collection.docx"):
    """
    Сохраняет структуру музыкальной коллекции в файл .docx
    :param music_library: Словарь с музыкальной коллекцией
    :param output_path: Путь для сохранения файла
    """
    try:
        # Создаем новый документ
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Добавляем заголовок
        title = doc.add_heading('Моя музыкальная коллекция', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Рекурсивно добавляем содержимое
        def add_items(node, level=1):
            for name, content in node.items():
                if name == "_files":
                    for file_name, _ in content:
                        p = doc.add_paragraph('    ' * level + f"🎵 {file_name}")
                        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
                else:
                    heading = doc.add_heading('    ' * (level - 1) + f"📁 {name}", level=min(level + 1, 6))
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    add_items(content, level + 1)

        add_items(music_library)

        # Сохраняем документ
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Ошибка при сохранении: {str(e)}")
        return False

# =================================================================



if __name__ == "__main__":
    root = tk.Tk()
    app = MusicCollectionApp(root)
    root.mainloop()