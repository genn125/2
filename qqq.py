
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




def save_to_docx(music_library, output_path="music_collection.docx"):
    """
    Сохраняет структуру музыкальной коллекции в файл .docx
    :param music_library: Словарь с музыкальной коллекцией
    :param output_path: Путь для сохранения файла
    """
    try:
        # Создаем новый документ
        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        # Добавляем заголовок
        title = doc.add_heading('Моя музыкальная коллекция', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Рекурсивно добавляем содержимое
        def add_items(node, level=1):
            for name, content in node.items():
                if name == "_files":
                    for file_name, _ in content:
                        p = doc.add_paragraph('    ' * level + f"🎵 {file_name}")
                        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
                else:
                    heading = doc.add_heading('    ' * (level - 1) + f"📁 {name}", level=min(level + 1, 6))
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    add_items(content, level + 1)

        add_items(music_library)

        # Сохраняем документ
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Ошибка при сохранении: {str(e)}")
        return False

def export_to_docx(self):
    """Экспорт коллекции в DOCX файл"""
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="Сохранить коллекцию как"
    )
    if file_path:
        if save_to_docx(self.music_library, file_path):
            self.status_bar.config(text=f"Коллекция сохранена: {file_path}", fg="green")
        else:
            self.status_bar.config(text="Ошибка при сохранении файла", fg="red")