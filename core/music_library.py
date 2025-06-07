import hashlib
import json
import os
import tkinter as tk
from tkinter import filedialog, ttk
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


def get_file_hash(filepath): # вычисляет MD5-хеш файла
    if not os.path.isfile(filepath):
        return None
    try:
        hasher = hashlib.md5()
        with open(filepath, 'rb') as f:
            buf = f.read(65536) # Читаем блоками по 64KB
            while len(buf) > 0:
                hasher.update(buf) # Добавляем данные в хеш
                buf = f.read(65536)  # Следующий блок
        return hasher.hexdigest()
    except (PermissionError, OSError):
        return None

def toggle_all(state, format_vars):
    for fmt, var in format_vars.items():
        var.set(1 if state else 0)

class MusicLibrary:
    def __init__(self):
        self.config_file = "config.json"
        self.all_formats = {
            'Аудио': ['.mp3', '.flac', '.wav', '.ogg', '.m4a', '.aac', '.wma'],
            'Изображения': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.raw', '.ico', '.tiff'],
            'Видео': ['.mp4', '.mpg', '.avi', '.mkv', '.flv', '.3gp', '.mov',' .wmv'],
            'Документы': ['.txt', 'pdf', '.rtf', '.docx', '.doc', '.xlsx', '.xls'],
            'Исполняемые': ['.exe', '.msi','.iso','.zip','.rar', '.7z'],
            'Разные': ['.dll', '.py', '.pyp', '.inf', '.js', '.css']
        }
        self.supported_formats = []
        self.music_library = defaultdict(lambda: defaultdict(dict))
        self.previous_files = set()
        self.scanning = False
        self.cancel_scan = False
        self.load_config()
        self.save_config()

    def load_config(self):
        try:
            with open(self.config_file, 'r') as f:
                config = json.load(f)
                self.supported_formats = config.get('formats', [])
        except (FileNotFoundError, json.JSONDecodeError):
            self.supported_formats = [fmt for group in self.all_formats.values() for fmt in group]

    def save_config(self):
        with open(self.config_file, 'w') as f:
            json.dump({'formats': self.supported_formats}, f)

    def show_format_selection(self, parent):
        """Окно выбора форматов с прогресс-баром"""
        selection_window = tk.Toplevel(parent)
        selection_window.title("Выбор формата файлов")
        selection_window.geometry("600x550")

        main_frame = tk.Frame(selection_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
# Создаем контейнер для колонок групп
        columns_frame = tk.Frame(main_frame)
        columns_frame.pack(fill=tk.BOTH, expand=True)
# Создаем 2 колонки
        left_column = tk.Frame(columns_frame)
        left_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        right_column = tk.Frame(columns_frame)
        right_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
# Скролл бар
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas)
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 10), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

# Общие настройки стиля кнопок
        main_button_style = {
            'font': ('Arial', 8, 'bold'),
            'width': 12,  # Устанавливаем одинаковую ширину width для всех кнопок
            'padx': 5,
            'pady': 2
        }
# Кнопки управления выбором
        select_all_frame = tk.Frame(scrollable_frame)
        select_all_frame.pack(fill=tk.X, pady=5)
        tk.Button(select_all_frame, text="Выбрать все", **main_button_style, bg="#4CAF50", fg="white",
                 command=lambda: toggle_all(True, format_vars)).pack(side=tk.LEFT, padx=2)
        tk.Button(select_all_frame, text="Снять все", **main_button_style,
                  command=lambda: toggle_all(False, format_vars)).pack(side=tk.LEFT)
        format_vars = {}
# Распределяем группы по колонкам
        group_frames = []
        for i, (group_name, formats) in enumerate(self.all_formats.items()):
# Выбираем колонку (четные - левая, нечетные - правая)
            column = left_column if i % 2 == 0 else right_column

            group_frame = tk.LabelFrame(column, text=group_name, padx=5, pady=5)
            group_frame.pack(fill=tk.X, pady=5)
            group_frames.append(group_frame)

            # Кнопки управления для группы
            group_buttons_frame = tk.Frame(group_frame)
            group_buttons_frame.pack(fill=tk.X, pady=2)

            def create_group_toggle(group_formats, state):
                return lambda: [format_vars[fmt].set(1 if state else 0) for fmt in group_formats]

            tk.Button(group_buttons_frame, text="✓",
                      command=create_group_toggle(formats, True),
                      width=3).pack(side=tk.LEFT, padx=2)
            tk.Button(group_buttons_frame, text="✕",
                      command=create_group_toggle(formats, False),
                      width=3).pack(side=tk.LEFT, padx=2)

            # 2 колонки внутри каждой группы для форматов
            cols = [tk.Frame(group_frame) for _ in range(2)]
            for col in cols:
                col.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

            for j, fmt in enumerate(formats):
                col = cols[j % 2]  # Распределение по 2 колонкам внутри группы
                format_vars[fmt] = tk.IntVar(value=1 if fmt in self.supported_formats else 0)
                tk.Checkbutton(col, text=fmt, variable=format_vars[fmt], anchor='w').pack(anchor='w')

        # Кнопки подтверждения
        button_frame = tk.Frame(main_frame)
        button_frame.pack(fill=tk.NONE, padx=10, pady=10)

        def apply_selection():
            self.supported_formats = [fmt for fmt, var in format_vars.items() if var.get()]
            self.save_config()
            selection_window.destroy()

        tk.Button(button_frame, text=" Применить ", **main_button_style, bg="#4CAF50", fg="white",
                  command=apply_selection).pack(side=tk.RIGHT)
        tk.Button(button_frame, text=" Отмена ", **main_button_style,
                  command=selection_window.destroy).pack(side=tk.LEFT, padx=2)

        selection_window.grab_set()
        selection_window.wait_window()
        return len(self.supported_formats) > 0

    def scan_folder(self, folder_path, clear_existing=True, parent=None):
        """Оптимизированное сканирование с прогресс-баром"""
        if not folder_path:
            return False, "Не выбрана папка"
# Окно прогресса
        progress = tk.Toplevel(parent) # Toplevel — это окно верхнего уровня, появляется поверх основного (parent).
        progress.title("Сканирование") #  заголовок окна
        progress.geometry("500x80")
#  Добавление текстовой метки
        tk.Label(progress, text=f"Идет сканирование папки {folder_path}").pack(pady=10) #  метка с отступом сверху и
        # снизу (10)
        '''
            ttk.Progressbar — прогресс-бар из модуля ttk (более современный, чем tkinter.Progressbar).
            mode="indeterminate" — анимированная полоса, которая движется бесконечно
            (подходит для процессов без известного времени выполнения).
        '''
        progress_bar = ttk.Progressbar(progress, mode="indeterminate")
        progress_bar.pack(pady=10)
        progress_bar.start()
        '''
            update() принудительно обновляет окно, чтобы оно отобразилось сразу, а не ждало завершения цикла tkinter.
            Без этого окно может "зависнуть" до окончания сканирования.
        '''
        progress.update()
        try:
            if clear_existing: # флаг, указывающий, нужно ли очистить текущую библиотеку перед сканированием.
                self.music_library.clear()
                self.music_library = defaultdict(lambda: defaultdict(dict)) # создаёт вложенную структуру словарей
                # для хранения данных:
            file_count = defaultdict(int) # Счётчик файлов по расширениям (например: {".mp3": 50})
            current_files = set() # Множество хешей уже обработанных файлов (для проверки новых)

            def scan_recursive(current_path, node):
                nonlocal file_count
                has_valid_files = False # флаг, указывающий, что в папке есть (нет) поддерживаемые файлы.
                try:
                    entries = os.listdir(current_path) # получает список файлов и папок в current_path.
                    progress.update()  # обновляет окно прогресса при каждом заходе в новую папку
                except (PermissionError, OSError):
                    return False
                for entry in entries: # Если entry — папка, функция вызывает саму себя для её сканирования.
                    try:
                        full_path = os.path.abspath(os.path.join(current_path, entry))
                        if os.path.isdir(full_path):
                            subnode = {}
                            if scan_recursive(full_path, subnode):
                                node[entry] = subnode
                                has_valid_files = True # флаг, указывающий, что в папке есть поддерживаемые файлы.
                        elif any(entry.lower().endswith(fmt) for fmt in self.supported_formats):#список поддерживаемых форматов
                            file_hash = get_file_hash(full_path)# вычисляет хеш (например, md5) для отслеживания
                            # изменений
                            if "_files" not in node:
                                node["_files"] = []
                            is_new = file_hash not in self.previous_files if file_hash else True
                            node["_files"].append((entry, full_path, is_new))# сохраняет имя, путь и флаг "новый файл"
                            ext = os.path.splitext(entry)[1].lower()
                            file_count[ext] += 1 # увеличивает счётчик для данного расширения
                            if file_hash:
                                current_files.add(file_hash) # добавляет хеш в множество обработанных файлов.
                            has_valid_files = True
                    except (PermissionError, OSError):
                        continue
                return has_valid_files
            success = scan_recursive(folder_path, self.music_library)
            self.previous_files = current_files
            if not success:
                return False, "Не найдено файлов с выбранными форматами"
            stats = ", ".join([f"{ext.upper()}: {count}" for ext, count in file_count.items()])
            return True, f"Найдено: {sum(file_count.values())} файлов ({stats})"
        finally:
            progress.destroy()  # Закрываем окно прогресса в любом случае

    def get_library(self):
        return self.music_library

    def clear_library(self):
        self.music_library.clear()
        self.music_library = defaultdict(lambda: defaultdict(dict))

    def save_to_docx(self, output_path=None):
        if not output_path:
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                title="Сохранить коллекцию как"
            )
            if not output_path:
                return False, ""
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            title = doc.add_heading('Моя музыкальная коллекция', level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            def add_items(node, level=1):
                for name, content in node.items():
                    if name == "_files":
                        for file_name, _, is_new in content:
                            p = doc.add_paragraph('    ' * level + f" {file_name}")
                            p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        heading = doc.add_heading('    ' * (level - 1) + f" {name}", level=min(level + 1, 6))
                        heading.runs[0].font.color.rgb = RGBColor(0, 0, 128)
                        add_items(content, level + 1)
            add_items(self.music_library)
            doc.save(output_path)
            return True, output_path
        except Exception as e:
            return False, str(e)